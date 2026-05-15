"""
Bulletin Generator
Opens the mededelingen Word template and replaces dynamic sections
with data from Dropbox, Scipio, and GKIN preekroster scrapers.
Static sections (Title, Jaarthema, Overdenking, Collecte informatie,
OLE info, Vakantieplanning, Bereikbaarheid) are kept as-is.
"""

import os
import copy
from datetime import datetime, timedelta
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from config import Config

TEMPLATE_PATH = os.path.join(Config.TEMPLATE_DIR, 'mededelingen_template.docx')
if not os.path.exists(TEMPLATE_PATH):
    TEMPLATE_PATH = os.path.join('./doc_templates', 'mededelingen_template.docx')


class BulletinGenerator:
    def __init__(self):
        self.output_dir = Config.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, mededelingen_date: datetime,
                 takenrooster_entry: Dict[str, Any],
                 takenrooster_entries: List[Dict[str, Any]],
                 mededelingen_data: Dict[str, Any],
                 birthday_data: Dict[str, Any],
                 preekroster_data: Dict[str, Any],
                 user_data: Dict[str, Any] = None) -> str:
        """Generate the bulletin Word document.

        Args:
            mededelingen_date: selected bulletin date
            takenrooster_entry: dict with 'predikant', 'ovd', 'opmerking'
            takenrooster_entries: all entries for next-7-day lookup
            mededelingen_data: from DropboxExcelReader.get_mededelingen()
            birthday_data: from ScipioScraper.get_birthday_list()
            preekroster_data: from PreekrosterScraper.get_preekroster()
            user_data: optional manually entered fields from the web form:
                overdenking_predikant, overdenking_thema, overdenking_schriftlezing,
                overdenking_content,
                collecte_contant, collecte_bonnen, collecte_bank, collecte_tikkie,
                collecte_ole, bezoekers_volwassenen, bezoekers_kinderen,
                dankoffer_url, dankoffer_qr_path,
                ole_url, ole_qr_path,
                activiteiten (list of {datum, activiteit, locatie})

        Returns:
            path to the generated .docx file.
        """
        if user_data is None:
            user_data = {}

        doc = Document(TEMPLATE_PATH)

        predikant = takenrooster_entry.get('predikant', '')
        ovd = takenrooster_entry.get('ovd', '')
        opmerking = takenrooster_entry.get('opmerking', '')

        # --- Dynamic header: update date in title ---
        self._update_title_date(doc, mededelingen_date)

        # --- Overdenking --- only update if at least one field is filled
        _ov_fields = ['overdenking_predikant', 'overdenking_thema',
                      'overdenking_schriftlezing', 'overdenking_content']
        if any(user_data.get(f, '').strip() for f in _ov_fields):
            self._update_overdenking(doc, user_data)

        # --- Dynamic welkomstwoord ---
        is_ole = 'OLE' in opmerking.upper()
        self._update_welkom_online(doc, is_ole)
        self._update_welkomstwoord(doc, mededelingen_date, predikant, ovd)
        self._update_aanstaande(doc, mededelingen_date, takenrooster_entries)

        # --- Update collecte table with dynamic dates and user-supplied amounts ---
        self._update_collecte_table(doc, mededelingen_date, user_data)

        # --- Update collecte informatie: dates, URLs, QR images ---
        self._update_collecte_informatie(doc, mededelingen_date, user_data)

        # --- Update activiteiten kalender ---
        self._update_activiteiten_table(doc, user_data.get('activiteiten', []))

        # --- Replace tables first (before paragraph removal shifts indices) ---
        birthday_table_data = birthday_data.get('birthday_table', [])
        self._replace_birthday_table(doc, birthday_table_data)

        am_table = preekroster_data.get('am_table', [])
        self._replace_preekroster_am(doc, am_table)

        ole_table = preekroster_data.get('ole_table', [])
        self._replace_preekroster_ole(doc, ole_table)

        # --- Replace text sections (Landelijke first to avoid index shift) ---
        self._replace_section(doc, 'Landelijke Mededelingen',
                              mededelingen_data.get('landelijke_nl', ''))
        self._replace_section(doc, 'Regionale mededelingen',
                              mededelingen_data.get('regionale_nl', ''))

        # --- Save ---
        filename = f'GKIN_Amstelveen_Mededelingen_{mededelingen_date.strftime("%y%m%d")}.docx'
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        print(f"Bulletin saved: {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Dynamic header/welkomstwoord helpers
    # ------------------------------------------------------------------

    def _format_dutch_date(self, dt: datetime) -> str:
        """Format date as Dutch string, e.g. '17 mei 2026'."""
        months = ['januari', 'februari', 'maart', 'april', 'mei', 'juni',
                  'juli', 'augustus', 'september', 'oktober', 'november', 'december']
        return f"{dt.day} {months[dt.month - 1]} {dt.year}"

    def _update_title_date(self, doc: Document, mededelingen_date: datetime):
        """Replace the date in the title (Heading 1) paragraph, preserving the image run."""
        p = doc.paragraphs[0]
        date_str = self._format_dutch_date(mededelingen_date)
        new_text = f"Mededelingen GKIN Amstelveen – {date_str}"

        # Preserve runs that contain drawings (images), clear the rest
        first_text_run = True
        for run in p.runs:
            has_drawing = bool(run._element.findall(qn('w:drawing')))
            if has_drawing:
                continue  # keep image run untouched
            if first_text_run:
                run.text = new_text
                first_text_run = False
            else:
                run.text = ''

    def _update_welkomstwoord(self, doc: Document, mededelingen_date: datetime,
                              predikant: str, ovd: str):
        """Update P19 with dynamic date, predikant (bold), and OvD (bold)."""
        # Find paragraph 19 — "Vandaag, zondag ..."
        target_para = None
        for p in doc.paragraphs:
            if p.text.strip().startswith('Vandaag,'):
                target_para = p
                break

        if target_para is None:
            print("Warning: Welkomstwoord 'Vandaag, zondag...' paragraph not found")
            return

        # Clear all runs
        p_element = target_para._element
        for run in list(p_element.findall(qn('w:r'))):
            p_element.remove(run)

        date_str = self._format_dutch_date(mededelingen_date)
        dutch_days = ['maandag', 'dinsdag', 'woensdag', 'donderdag',
                      'vrijdag', 'zaterdag', 'zondag']
        day_name = dutch_days[mededelingen_date.weekday()]

        # Build new runs: "Vandaag, {dag} {date}, gaat voor " + bold(predikant) + ...
        target_para.add_run(f"Vandaag, {day_name} {date_str}, gaat voor ")
        run_pred = target_para.add_run(predikant)
        run_pred.bold = True
        target_para.add_run(". De ouderling van dienst is ")
        run_ovd = target_para.add_run(ovd)
        run_ovd.bold = True
        target_para.add_run(". Als u vragen heeft, kunt u de ouderling van dienst aanspreken.")

    def _update_overdenking(self, doc: Document, user_data: Dict[str, Any]):
        """Update the Overdenking heading and paragraphs with user-supplied content."""
        pred = user_data.get('overdenking_predikant', '').strip()
        thema = user_data.get('overdenking_thema', '').strip()
        schrift = user_data.get('overdenking_schriftlezing', '').strip()
        content = user_data.get('overdenking_content', '').strip()

        # Update heading "Overdenking door ..."
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading') and 'overdenking' in p.text.lower():
                for run in p.runs:
                    if not run._element.findall(qn('w:drawing')):
                        run.text = ''
                if p.runs:
                    p.runs[0].text = f"Overdenking door {pred}" if pred else "Overdenking"
                break

        # Update P7-type: Thema + Schriftlezing (first Normal after heading)
        in_overdenking = False
        thema_done = False
        content_paras = []
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading') and 'overdenking' in p.text.lower():
                in_overdenking = True
                continue
            if in_overdenking and p.style.name.startswith('Heading'):
                break
            if in_overdenking and p.style.name == 'Normal':
                if not thema_done:
                    # Rewrite thema/schriftlezing line
                    p_el = p._element
                    for r in list(p_el.findall(qn('w:r'))):
                        p_el.remove(r)
                    if thema or schrift:
                        r1 = p.add_run('Thema: ')
                        r1.bold = True
                        r2 = p.add_run(f'"{thema}"' if thema else '')
                        r2.bold = True
                        if schrift:
                            p.add_run('\n')
                            r3 = p.add_run('Schriftlezing')
                            r3.bold = True
                            p.add_run(f': {schrift}')
                    thema_done = True
                else:
                    content_paras.append(p)

        # Replace content paragraphs with user-supplied text.
        # Collapse to one paragraph per non-empty line (no double-blank gaps),
        # left-aligned. Choose font size based on estimated line count to fit column.
        # Column fits ~28 lines at 10pt; scale down only if content exceeds that.
        _COL_WIDTH_CHARS = 80   # approx chars per line at 10pt in column
        _lines_raw = [l.strip() for l in content.replace('\r', '').split('\n') if l.strip()] if content else []
        # Count wrapped lines
        _total_lines = sum(max(1, (len(l) + _COL_WIDTH_CHARS - 1) // _COL_WIDTH_CHARS) for l in _lines_raw)
        # Column fits ~28 lines at 10pt before spilling to right column
        if _total_lines <= 28:
            CONTENT_PT = 10
        elif _total_lines <= 32:
            CONTENT_PT = 9.5
        else:
            CONTENT_PT = 9

        def _make_content_run(p_el, text: str):
            """Add a run with content font size set."""
            run_el = p_el.makeelement(qn('w:r'), {})
            rPr = run_el.makeelement(qn('w:rPr'), {})
            sz = rPr.makeelement(qn('w:sz'), {})
            sz.set(qn('w:val'), str(int(CONTENT_PT * 2)))  # half-points
            szCs = rPr.makeelement(qn('w:szCs'), {})
            szCs.set(qn('w:val'), str(int(CONTENT_PT * 2)))
            rPr.append(sz)
            rPr.append(szCs)
            run_el.append(rPr)
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = text
            t_el.set(qn('xml:space'), 'preserve')
            run_el.append(t_el)
            p_el.append(run_el)

        def _set_para_left(p_el):
            """Force left alignment and remove space-before/after on a paragraph."""
            pPr = p_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = p_el.makeelement(qn('w:pPr'), {})
                p_el.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = pPr.makeelement(qn('w:jc'), {})
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
            # Remove paragraph spacing
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = pPr.makeelement(qn('w:spacing'), {})
                pPr.append(spacing)
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')

        if content and content_paras:
            # Flatten: split on any newline, discard empty lines, one line per para
            lines = [l.strip() for l in content.replace('\r', '').split('\n') if l.strip()]

            ref_el = content_paras[0]._element

            # Reuse existing paragraphs first, then clone, then remove leftover
            all_els = [p._element for p in content_paras]

            # We need exactly len(lines) paragraphs
            while len(all_els) < len(lines):
                new_p = copy.deepcopy(ref_el)
                all_els[-1].addnext(new_p)
                all_els.append(new_p)

            # Remove surplus paragraphs
            for surplus in all_els[len(lines):]:
                surplus.getparent().remove(surplus)
            all_els = all_els[:len(lines)]

            for p_el, line in zip(all_els, lines):
                # Clear existing runs
                for r in list(p_el.findall(qn('w:r'))):
                    p_el.remove(r)
                _set_para_left(p_el)
                _make_content_run(p_el, line)

    def _update_welkom_online(self, doc: Document, is_ole: bool):
        """Update P18: include 'Online ' before 'Eredienst' only if OLE."""
        target_para = None
        for p in doc.paragraphs:
            if 'van harte welkom bij deze' in p.text:
                target_para = p
                break

        if target_para is None:
            return

        # Rebuild text: replace "Online Eredienst" or "Eredienst"
        full_text = target_para.text
        if is_ole:
            new_text = full_text  # keep "Online" as-is in template
        else:
            new_text = full_text.replace('Online Eredienst', 'Eredienst')

        # Only rewrite if changed
        if new_text != full_text:
            p_element = target_para._element
            for run in list(p_element.findall(qn('w:r'))):
                p_element.remove(run)
            target_para.add_run(new_text)

    def _update_aanstaande(self, doc: Document, mededelingen_date: datetime,
                           takenrooster_entries: List[Dict[str, Any]]):
        """Replace 'Aanstaande...' paragraphs with upcoming services in next 7 days."""
        dutch_days = ['maandag', 'dinsdag', 'woensdag', 'donderdag',
                      'vrijdag', 'zaterdag', 'zondag']

        # Find upcoming entries: date > selected and date <= selected + 7
        upcoming = []
        for e in takenrooster_entries:
            d = e['date']
            if isinstance(d, datetime):
                d_date = d.date()
            else:
                d_date = d
            sel_date = mededelingen_date.date() if isinstance(mededelingen_date, datetime) else mededelingen_date
            if sel_date < d_date <= sel_date + timedelta(days=7):
                upcoming.append(e)
        upcoming.sort(key=lambda e: e['date'])

        # Find existing "Aanstaande..." paragraphs
        aanstaande_paras = []
        for p in doc.paragraphs:
            if p.text.strip().startswith('Aanstaande'):
                aanstaande_paras.append(p)

        # Remove existing "Aanstaande..." paragraphs
        for p in aanstaande_paras:
            p._element.getparent().remove(p._element)

        # Insert new "Aanstaande..." paragraphs before the next element after welkomstwoord
        # Find the "Vandaag, zondag..." paragraph and insert after it
        vandaag_el = None
        for p in doc.paragraphs:
            if p.text.strip().startswith('Vandaag, zondag') or p.text.strip().startswith('Vandaag,  zondag'):
                vandaag_el = p._element
                break

        if vandaag_el is None:
            return

        # Insert new paragraphs after vandaag_el
        insert_after = vandaag_el
        for e in upcoming:
            d = e['date']
            day_name = dutch_days[d.weekday()]
            date_str = self._format_dutch_date(d)
            predikant = e.get('predikant', '')
            opmerking = e.get('opmerking', '')

            # Extract dienst type from opmerking
            dienst_type = self._extract_dienst_type(opmerking)
            if dienst_type:
                dienst_text = f"{dienst_type} "
            else:
                dienst_text = ""

            # Create paragraph
            new_p = vandaag_el.getparent().makeelement(qn('w:p'), {})
            # Copy paragraph properties (style) from vandaag
            pPr_src = vandaag_el.find(qn('w:pPr'))
            if pPr_src is not None:
                new_p.append(copy.deepcopy(pPr_src))

            # Build runs
            self._add_run_to_element(new_p, f"Aanstaande {day_name} {date_str}, hoopt in de {dienst_text}Eredienst in Amstelveen voor te gaan, ")
            self._add_run_to_element(new_p, predikant, bold=True)
            self._add_run_to_element(new_p, ". Aanvang is om 10:30 uur. ")

            insert_after.addnext(new_p)
            insert_after = new_p

    def _extract_dienst_type(self, opmerking: str) -> str:
        """Extract service type from opmerking, e.g. 'Hemelvaart', 'Pinksteren'."""
        if not opmerking:
            return ''
        # Strip OLE link part
        text = opmerking.split('OLE')[0].strip().rstrip(',').strip()
        # Known types
        if text:
            return text
        if 'OLE' in opmerking.upper():
            return 'OLE'
        return ''

    def _set_amount_cell(self, cell, amount: str):
        """Write an amount cell with € left and amount right-aligned via tab stop.

        Result: '€<TAB>181,35' with a right tab stop at the cell width,
        so all decimal points align on the right edge.
        """
        # Get cell width in twips from tcPr/tcW
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        cell_w = 1985  # fallback twips (~3.5cm)
        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                try:
                    cell_w = int(tcW.get(qn('w:w'), cell_w))
                except (ValueError, TypeError):
                    pass

        for para in cell.paragraphs:
            p_el = para._element
            # Clear existing runs
            for r in list(p_el.findall(qn('w:r'))):
                p_el.remove(r)

            # Add/update pPr with a right tab stop at cell width
            pPr = p_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = p_el.makeelement(qn('w:pPr'), {})
                p_el.insert(0, pPr)
            # Remove old tabs
            old_tabs = pPr.find(qn('w:tabs'))
            if old_tabs is not None:
                pPr.remove(old_tabs)
            tabs_el = pPr.makeelement(qn('w:tabs'), {})
            tab_el = tabs_el.makeelement(qn('w:tab'), {})
            tab_el.set(qn('w:val'), 'right')
            tab_el.set(qn('w:pos'), str(cell_w))
            tabs_el.append(tab_el)
            pPr.append(tabs_el)

            # Run 1: '€'
            r1 = p_el.makeelement(qn('w:r'), {})
            t1 = r1.makeelement(qn('w:t'), {})
            t1.text = '€'
            r1.append(t1)
            p_el.append(r1)

            # Run 2: tab character
            r_tab = p_el.makeelement(qn('w:r'), {})
            tab_run = r_tab.makeelement(qn('w:tab'), {})
            r_tab.append(tab_run)
            p_el.append(r_tab)

            # Run 3: amount
            r2 = p_el.makeelement(qn('w:r'), {})
            t2 = r2.makeelement(qn('w:t'), {})
            t2.text = amount if amount else ''
            t2.set(qn('xml:space'), 'preserve')
            r2.append(t2)
            p_el.append(r2)
            break  # only first paragraph

    def _fill_collecte_table(self, table, entry: Dict[str, Any]):
        """Fill a single collecte table from one opbrengst entry dict."""
        service_date = entry.get('service_date', '')

        # R0: header
        if service_date:
            header_text = f"GKIN Amstelveen {service_date}"
            for c in range(2):
                self._set_cell_text_preserving(table.cell(0, c), header_text)

        # R2-R5: amounts
        amounts = [
            entry.get('collecte_contant', ''),
            entry.get('collecte_bonnen', ''),
            entry.get('collecte_bank', ''),
            entry.get('collecte_tikkie', ''),
        ]
        for i, r in enumerate([2, 3, 4, 5]):
            val = amounts[i].strip() if amounts[i] else ''
            self._set_amount_cell(table.cell(r, 1), val)

        # R6: Totaal
        try:
            total = sum(float(a.replace(',', '.').strip()) for a in amounts if a.strip())
            total_str = f'{total:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
        except Exception:
            total_str = ''
        self._set_amount_cell(table.cell(6, 1), total_str)

        # R7: Bezoekers
        vw = entry.get('bezoekers_volwassenen', '').strip()
        ki = entry.get('bezoekers_kinderen', '').strip()
        if vw or ki:
            try:
                totaal = str(int(vw or 0) + int(ki or 0))
            except Exception:
                totaal = ''
            bezoekers_text = f"Aantal bezoekers: {vw} volwassenen, {ki} kinderen, {totaal} totaal bezoekers"
        else:
            bezoekers_text = "Aantal bezoekers:    volwassenen,    kinderen,    totaal bezoekers"
        for c in range(2):
            self._set_cell_text_preserving(table.cell(7, c), bezoekers_text)

        # Extra items (HSK, bijbelstudie, etc.) — insert rows before OLE row
        extra_items = entry.get('extra_items', [])
        tbl_el = table._tbl
        rows = tbl_el.findall(qn('w:tr'))
        ole_row_el = rows[9]  # R9 is the OLE row

        # Remove any previously inserted extra rows (marked with custom attr)
        for r in list(tbl_el.findall(qn('w:tr'))):
            if r.get('w:customExtra') == '1':
                tbl_el.remove(r)

        def _set_tc_text(tc_el, text: str):
            """Clear all runs in a table cell and write a single clean run."""
            for p_el in tc_el.findall(qn('w:p')):
                for r in list(p_el.findall(qn('w:r'))):
                    p_el.remove(r)
                # Write one run with the text
                run_el = p_el.makeelement(qn('w:r'), {})
                t_el = run_el.makeelement(qn('w:t'), {})
                t_el.text = text
                t_el.set(qn('xml:space'), 'preserve')
                run_el.append(t_el)
                p_el.append(run_el)
                break  # only write to the first paragraph in the cell

        def _set_amount_tc(tc_el, amount: str, cell_w: int = 1985):
            """Like _set_amount_cell but operates directly on a w:tc element."""
            for p_el in tc_el.findall(qn('w:p')):
                for r in list(p_el.findall(qn('w:r'))): p_el.remove(r)
                pPr = p_el.find(qn('w:pPr'))
                if pPr is None:
                    pPr = p_el.makeelement(qn('w:pPr'), {}); p_el.insert(0, pPr)
                old = pPr.find(qn('w:tabs'))
                if old is not None: pPr.remove(old)
                tabs = pPr.makeelement(qn('w:tabs'), {})
                tab = tabs.makeelement(qn('w:tab'), {})
                tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), str(cell_w))
                tabs.append(tab); pPr.append(tabs)
                r1 = p_el.makeelement(qn('w:r'), {})
                t1 = r1.makeelement(qn('w:t'), {}); t1.text = '€'; r1.append(t1); p_el.append(r1)
                rt = p_el.makeelement(qn('w:r'), {}); rt.append(p_el.makeelement(qn('w:tab'), {})); p_el.append(rt)
                r2 = p_el.makeelement(qn('w:r'), {})
                t2 = r2.makeelement(qn('w:t'), {}); t2.text = amount; t2.set(qn('xml:space'), 'preserve'); r2.append(t2); p_el.append(r2)
                break

        for item in extra_items:
            new_row = copy.deepcopy(ole_row_el)
            new_row.set('w:customExtra', '1')
            cells_in_row = new_row.findall(qn('w:tc'))
            if len(cells_in_row) >= 2:
                _set_tc_text(cells_in_row[0], item.get('desc', ''))
                _set_amount_tc(cells_in_row[1], item.get('amount', ''))
            ole_row_el.addprevious(new_row)

        # Update OLE row (always last row)
        rows_updated = tbl_el.findall(qn('w:tr'))
        ole_row = rows_updated[-1]
        ole_cells = ole_row.findall(qn('w:tc'))
        ole_val = entry.get('collecte_ole', '').strip()
        if len(ole_cells) >= 2:
            _set_tc_text(ole_cells[0], f"Collecte opbrengst OLE {service_date}")
            _set_amount_tc(ole_cells[1], ole_val)

    def _update_collecte_table(self, doc: Document, mededelingen_date: datetime,
                               user_data: Dict[str, Any] = None):
        """Update collecte table(s) with dynamic dates and user amounts.

        If user_data contains 'opbrengst_entries' (list of per-service dicts from email),
        one table block is rendered per entry. Extra tables are cloned from the template.
        Falls back to single-table behaviour using flat user_data fields.
        """
        if user_data is None:
            user_data = {}

        dutch_days = ['maandag', 'dinsdag', 'woensdag', 'donderdag',
                      'vrijdag', 'zaterdag', 'zondag']
        prev_date     = mededelingen_date - timedelta(days=7)
        prev_date_str = self._format_dutch_date(prev_date)

        entries = user_data.get('opbrengst_entries', [])

        # Build a single fallback entry from flat user_data fields
        if not entries:
            prev_day_name = dutch_days[prev_date.weekday()]
            entries = [{
                'service_date': f"{prev_day_name} {prev_date_str}",
                'collecte_contant':      user_data.get('collecte_contant', ''),
                'collecte_bonnen':       user_data.get('collecte_bonnen', ''),
                'collecte_bank':         user_data.get('collecte_bank', ''),
                'collecte_tikkie':       user_data.get('collecte_tikkie', ''),
                'collecte_ole':          user_data.get('collecte_ole', ''),
                'bezoekers_volwassenen': user_data.get('bezoekers_volwassenen', ''),
                'bezoekers_kinderen':    user_data.get('bezoekers_kinderen', ''),
            }]

        base_table = doc.tables[0]

        # Fill the first (existing) table
        self._fill_collecte_table(base_table, entries[0])

        # For each additional entry, deep-clone the table and insert after the previous one
        prev_tbl_el = base_table._tbl
        for entry in entries[1:]:
            new_tbl = copy.deepcopy(base_table._tbl)
            prev_tbl_el.addnext(new_tbl)
            # Also insert a blank paragraph between tables for spacing
            spacer = OxmlElement('w:p')
            new_tbl.addprevious(spacer)
            # Wrap in a docx Table object to use _fill_collecte_table
            from docx.table import Table as DocxTable
            new_table_obj = DocxTable(new_tbl, doc)
            self._fill_collecte_table(new_table_obj, entry)
            prev_tbl_el = new_tbl

    def _update_collecte_informatie(self, doc: Document, mededelingen_date: datetime,
                                    user_data: Dict[str, Any] = None):
        """Update dates in Dankoffer and OLE sections, replace URLs and QR images."""
        if user_data is None:
            user_data = {}
        date_ddmmyyyy = mededelingen_date.strftime('%d-%m-%Y')
        dankoffer_url = user_data.get('dankoffer_url', '').strip() or '[Insert URL here]'
        ole_url = user_data.get('ole_url', '').strip() or '[Insert URL here]'

        for p in doc.paragraphs:
            if 'NL40.ABNA.0549.3085.12' in p.text and 'Collecte' in p.text:
                p_element = p._element
                for run in list(p_element.findall(qn('w:r'))):
                    p_element.remove(run)
                p.add_run(
                    'Overmaken naar rekeningnummer van GKIN Amstelveen, '
                    f'IBAN: NL40.ABNA.0549.3085.12, o.v.v. \u201cCollecte {date_ddmmyyyy}\u201d'
                )

            if 'QR-code' in p.text and 'betaalverzoek' in p.text:
                # Clear all existing content (runs + hyperlinks), rewrite label + clickable URL
                p_element = p._element
                for child in list(p_element):
                    if child.tag in (qn('w:r'), qn('w:hyperlink')):
                        p_element.remove(child)
                p.add_run('Gebruik te maken van de QR-code, of betaalverzoek link: ')
                if dankoffer_url and dankoffer_url != '[Insert URL here]':
                    self._add_hyperlink(p, dankoffer_url)
                else:
                    p.add_run(dankoffer_url)

            if 'landelijke kas' in p.text and 'INGB' in p.text:
                p_element = p._element
                for hl in list(p_element.findall(qn('w:hyperlink'))):
                    p_element.remove(hl)
                for run in list(p_element.findall(qn('w:r'))):
                    p_element.remove(run)
                p.add_run(
                    'Het dankoffer voor de landelijke kas kunt u overmaken naar het landelijke '
                    'rekeningnummer van Gereja Kristen Indonesia Nederland: '
                    f'IBAN: NL19 INGB 000 261 8290, O.v.v. \u201cCollecte OLE {date_ddmmyyyy}\u201d '
                )
                if ole_url and ole_url != '[Insert URL here]':
                    self._add_hyperlink(p, ole_url)
                else:
                    p.add_run(ole_url)

            if not p.text.strip() and p._element.findall(qn('w:hyperlink')):
                self._replace_hyperlinks(p, '')

        # Replace QR image paragraphs
        dankoffer_qr_path = user_data.get('dankoffer_qr', '').strip()
        ole_qr_path = user_data.get('ole_qr', '').strip()
        qr_paths = iter([dankoffer_qr_path, ole_qr_path])

        in_collecte = False
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading') and 'collecte informatie' in p.text.strip().lower():
                in_collecte = True
                continue
            if p.style.name.startswith('Heading') and 'regionale' in p.text.strip().lower():
                break
            if in_collecte:
                drawings = p._element.findall('.//' + qn('w:drawing'))
                if drawings:
                    qr_path = next(qr_paths, '')
                    for run in list(p._element.findall(qn('w:r'))):
                        p._element.remove(run)
                    if qr_path and os.path.isfile(qr_path):
                        from docx.shared import Cm
                        p.add_run().add_picture(qr_path, width=Cm(3), height=Cm(3))
                    else:
                        p.add_run('[Insert QR here]')

    def _add_hyperlink(self, para, url: str, display_text: str = None):
        """Insert a clickable hyperlink run into a paragraph."""
        if not url:
            return
        # Register the relationship
        part = para.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        # Build w:hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        hyperlink.set(qn('w:history'), '1')
        # Build the run inside
        run_el = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        # Apply hyperlink character style (blue + underline)
        style_el = OxmlElement('w:rStyle')
        style_el.set(qn('w:val'), 'Hyperlink')
        rpr.append(style_el)
        run_el.append(rpr)
        t_el = OxmlElement('w:t')
        t_el.text = display_text or url
        t_el.set(qn('xml:space'), 'preserve')
        run_el.append(t_el)
        hyperlink.append(run_el)
        para._element.append(hyperlink)

    def _replace_hyperlinks(self, para, replacement_url: str):
        """Replace existing hyperlink elements in a paragraph with a new clickable hyperlink."""
        p_element = para._element
        for hl in list(p_element.findall(qn('w:hyperlink'))):
            p_element.remove(hl)
        if replacement_url:
            if para.runs:
                para.runs[-1].text = para.runs[-1].text.rstrip() + ' '
            self._add_hyperlink(para, replacement_url)

    def _update_activiteiten_table(self, doc: Document, activiteiten: List[Dict] = None):
        """Keep activiteiten table header, fill with user-supplied rows or 1 empty row."""
        table = self._find_table_after_heading(doc, 'Activiteiten kalender')
        if table is None:
            print("Warning: activiteiten table not found")
            return

        tbl_element = table._tbl
        rows = list(tbl_element.findall(qn('w:tr')))
        if len(rows) < 1:
            return

        header_row = rows[0]
        # Remove all existing data rows
        for tr in rows[1:]:
            tbl_element.remove(tr)

        # Determine rows to add
        if not activiteiten:
            activiteiten = [{'datum': '', 'activiteit': '', 'locatie': ''}]

        for act in activiteiten:
            new_row = copy.deepcopy(header_row)
            cells = new_row.findall(qn('w:tc'))

            # Column 0: datum [line-break] tijd
            col0_lines = [act.get('datum', '')]
            if act.get('tijd'):
                import re as _re
                _tijd = act['tijd'].strip()
                # If it looks like a time (HH:MM or HH.MM), format as Dutch "10.30u"
                _tm = _re.match(r'^(\d{1,2})[:.](\d{2})$', _tijd)
                if _tm:
                    col0_lines.append(f"{_tm.group(1)}.{_tm.group(2)}u")
                else:
                    # Free text like "na dienst", "TBV" — use as-is
                    col0_lines.append(_tijd)

            # Column 1: activiteit name [line-break] o.l.v. Persoon
            col1_lines = [act.get('activiteit', '')]
            if act.get('olv'):
                col1_lines.append(f"o.l.v. {act['olv']}")

            # Column 2: locatie
            col2_lines = [act.get('locatie', '')]

            all_cols = [col0_lines, col1_lines, col2_lines]

            for ci, tc in enumerate(cells[:3]):
                lines = all_cols[ci]
                for p in tc.findall(qn('w:p')):
                    for run in p.findall(qn('w:r')):
                        p.remove(run)
                    for li, line in enumerate(lines):
                        if not line:
                            continue
                        if li > 0:
                            # Add a line break run before next line
                            br_run = p.makeelement(qn('w:r'), {})
                            br_el = br_run.makeelement(qn('w:br'), {})
                            br_run.append(br_el)
                            p.append(br_run)
                        run_el = p.makeelement(qn('w:r'), {})
                        t_el = run_el.makeelement(qn('w:t'), {})
                        t_el.text = line
                        t_el.set(qn('xml:space'), 'preserve')
                        run_el.append(t_el)
                        p.append(run_el)
            tbl_element.append(new_row)

    def _set_cell_text_preserving(self, cell, text: str):
        """Set cell text while preserving cell formatting (borders, shading, etc)."""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = text
            else:
                paragraph.add_run(text)
            break  # only update first paragraph

    def _add_run_to_element(self, p_element, text: str, bold: bool = False):
        """Add a run with text to a paragraph element."""
        run_el = p_element.makeelement(qn('w:r'), {})
        if bold:
            rPr = run_el.makeelement(qn('w:rPr'), {})
            b_el = rPr.makeelement(qn('w:b'), {})
            rPr.append(b_el)
            run_el.append(rPr)
        t_el = run_el.makeelement(qn('w:t'), {})
        t_el.text = text
        t_el.set(qn('xml:space'), 'preserve')
        run_el.append(t_el)
        p_element.append(run_el)

    # ------------------------------------------------------------------
    # Section replacement helpers
    # ------------------------------------------------------------------

    def _find_heading_index(self, doc: Document, heading_text: str) -> int:
        """Find paragraph index whose text starts with heading_text (case-insensitive)."""
        target = heading_text.strip().lower()
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading') and p.text.strip().lower().startswith(target):
                return i
        return -1

    def _find_next_heading_index(self, doc: Document, start: int) -> int:
        """Find the index of the next Heading paragraph after start."""
        for i in range(start + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].style.name.startswith('Heading'):
                return i
        return len(doc.paragraphs)

    def _replace_section(self, doc: Document, heading_text: str, new_content: str):
        """Replace all Normal paragraphs between a heading and the next heading.

        Strategy: find heading element in body, collect all Normal paragraph
        elements between it and the next heading, remove them, then insert
        new paragraph elements in their place.
        """
        body = doc.element.body
        elem_to_para = {p._element: p for p in doc.paragraphs}
        target = heading_text.strip().lower()

        # Find the heading element
        heading_el = None
        for element in body:
            if not element.tag.endswith('}p'):
                continue
            para = elem_to_para.get(element)
            if para and para.style.name.startswith('Heading') and para.text.strip().lower().startswith(target):
                heading_el = element
                break

        if heading_el is None:
            print(f"Warning: heading '{heading_text}' not found in template")
            return

        # Collect Normal paragraph elements between heading and next heading/table
        content_elements = []
        next_anchor = None
        started = False
        for element in body:
            if element is heading_el:
                started = True
                continue
            if not started:
                continue
            if element.tag.endswith('}p'):
                para = elem_to_para.get(element)
                if para and para.style.name.startswith('Heading'):
                    next_anchor = element
                    break
                content_elements.append(element)
            elif element.tag.endswith('}tbl'):
                next_anchor = element
                break

        # Keep a reference Normal paragraph element for cloning style
        ref_el = content_elements[0] if content_elements else None

        # Remove old content paragraphs
        for el in content_elements:
            body.remove(el)

        # Split new content into paragraphs on double newline
        blocks = [b.strip() for b in new_content.split('\n\n') if b.strip()]

        # Insert new paragraphs: first line of each block = Heading 3, rest = Normal
        for block in blocks:
            lines = block.split('\n')
            title_line = lines[0]
            body_lines = lines[1:]

            # --- Title paragraph (Heading 3, left-aligned) ---
            title_p = self._make_paragraph(body, 'Heading3', title_line, left_align=True)
            if next_anchor is not None:
                next_anchor.addprevious(title_p)
            else:
                body.append(title_p)

            # --- Body paragraph (Normal, left-aligned) with line breaks ---
            if body_lines:
                if ref_el is not None:
                    body_p = copy.deepcopy(ref_el)
                    for child in list(body_p):
                        if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                            body_p.remove(child)
                    # Force left align on the copied paragraph
                    pPr = body_p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = body_p.makeelement(qn('w:pPr'), {})
                        body_p.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = pPr.makeelement(qn('w:jc'), {})
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'left')
                else:
                    body_p = self._make_paragraph(body, 'Normal', '', left_align=True)

                # Clear any existing runs (in case deepcopy brought some)
                for run in list(body_p.findall(qn('w:r'))):
                    body_p.remove(run)

                for li, line in enumerate(body_lines):
                    run_el = body_p.makeelement(qn('w:r'), {})
                    t_el = run_el.makeelement(qn('w:t'), {})
                    t_el.text = line
                    t_el.set(qn('xml:space'), 'preserve')
                    run_el.append(t_el)
                    body_p.append(run_el)
                    if li < len(body_lines) - 1:
                        br_run = body_p.makeelement(qn('w:r'), {})
                        br_el = br_run.makeelement(qn('w:br'), {})
                        br_run.append(br_el)
                        body_p.append(br_run)

                if next_anchor is not None:
                    next_anchor.addprevious(body_p)
                else:
                    body.append(body_p)

    def _make_paragraph(self, body, style_name: str, text: str, left_align: bool = False):
        """Create a new paragraph element with given style and text."""
        new_p = body.makeelement(qn('w:p'), {})
        pPr = new_p.makeelement(qn('w:pPr'), {})
        pStyle = pPr.makeelement(qn('w:pStyle'), {})
        pStyle.set(qn('w:val'), style_name)
        pPr.append(pStyle)
        if left_align:
            jc = pPr.makeelement(qn('w:jc'), {})
            jc.set(qn('w:val'), 'left')
            pPr.append(jc)
        new_p.append(pPr)

        if text:
            run_el = new_p.makeelement(qn('w:r'), {})
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = text
            t_el.set(qn('xml:space'), 'preserve')
            run_el.append(t_el)
            new_p.append(run_el)

        return new_p

    # ------------------------------------------------------------------
    # Table replacement helpers
    # ------------------------------------------------------------------

    def _clear_and_resize_table(self, table, num_rows: int):
        """Remove existing data rows and add the needed number of rows."""
        # Remove all existing rows
        tbl = table._tbl
        for tr in list(tbl.findall(qn('w:tr'))):
            tbl.remove(tr)

        # Add new empty rows
        for _ in range(num_rows):
            tr = copy.deepcopy(table.rows[0]._tr) if len(table.rows) > 0 else tbl.makeelement(qn('w:tr'), {})
            tbl.append(tr)

    def _replace_birthday_table(self, doc: Document, birthday_table: List[Dict]):
        """Replace the birthday table (2-column layout)."""
        # Find the birthday table — it's after "Verjaardagen" heading
        bday_table = self._find_table_after_heading(doc, 'Verjaardagen')
        if bday_table is None:
            print("Warning: birthday table not found")
            return

        tbl_element = bday_table._tbl
        # Keep reference row for formatting
        existing_rows = list(tbl_element.findall(qn('w:tr')))
        ref_row = copy.deepcopy(existing_rows[0]) if existing_rows else None

        # Clear all rows
        for tr in existing_rows:
            tbl_element.remove(tr)

        # Add new rows
        for row_data in birthday_table:
            new_row = copy.deepcopy(ref_row) if ref_row is not None else tbl_element.makeelement(qn('w:tr'), {})
            cells = new_row.findall(qn('w:tc'))
            if len(cells) >= 2:
                self._set_cell_text(cells[0], row_data.get('left', ''))
                self._set_cell_text(cells[1], row_data.get('right', ''))
            tbl_element.append(new_row)

    def _replace_preekroster_am(self, doc: Document, am_table: List[Dict]):
        """Replace the AM preekrooster table (4 columns: date, lang, predikant, time)."""
        table = self._find_table_after_heading(doc, 'GKIN Amstelveen')
        if table is None:
            print("Warning: AM preekrooster table not found")
            return

        tbl_element = table._tbl
        existing_rows = list(tbl_element.findall(qn('w:tr')))
        ref_row = copy.deepcopy(existing_rows[0]) if existing_rows else None

        for tr in existing_rows:
            tbl_element.remove(tr)

        for entry in am_table:
            new_row = copy.deepcopy(ref_row) if ref_row is not None else tbl_element.makeelement(qn('w:tr'), {})
            cells = new_row.findall(qn('w:tc'))
            values = [entry.get('date', ''), entry.get('language', ''),
                      entry.get('predikant', ''), entry.get('time', '')]
            for ci, val in enumerate(values):
                if ci < len(cells):
                    self._set_cell_text(cells[ci], val)
            tbl_element.append(new_row)

    def _replace_preekroster_ole(self, doc: Document, ole_table: List[Dict]):
        """Replace the OLE preekrooster table (5 columns: date, lang, regio, predikant, time)."""
        table = self._find_table_after_heading(doc, 'Online Landelijke Eredienst (OLE)')
        if table is None:
            print("Warning: OLE preekrooster table not found")
            return

        tbl_element = table._tbl
        existing_rows = list(tbl_element.findall(qn('w:tr')))
        ref_row = copy.deepcopy(existing_rows[0]) if existing_rows else None

        for tr in existing_rows:
            tbl_element.remove(tr)

        for entry in ole_table:
            new_row = copy.deepcopy(ref_row) if ref_row is not None else tbl_element.makeelement(qn('w:tr'), {})
            cells = new_row.findall(qn('w:tc'))
            values = [entry.get('date', ''), entry.get('language', ''),
                      entry.get('regio', ''), entry.get('predikant', ''),
                      entry.get('time', '')]
            for ci, val in enumerate(values):
                if ci < len(cells):
                    self._set_cell_text(cells[ci], val)
            tbl_element.append(new_row)

    def _find_table_after_heading(self, doc: Document, heading_text: str):
        """Find the first table that appears after a specific heading in the document body."""
        target = heading_text.strip().lower()
        found_heading = False

        # Build a lookup from element to paragraph for fast access
        elem_to_para = {p._element: p for p in doc.paragraphs}
        elem_to_table = {t._tbl: t for t in doc.tables}

        for element in doc.element.body:
            if element.tag.endswith('}p') and not found_heading:
                para = elem_to_para.get(element)
                if para and para.style.name.startswith('Heading') and para.text.strip().lower().startswith(target):
                    found_heading = True
            elif element.tag.endswith('}tbl') and found_heading:
                table = elem_to_table.get(element)
                if table:
                    return table
                break

        return None

    def _set_cell_text(self, tc_element, text: str):
        """Set the text content of a table cell, preserving formatting."""
        # Find all paragraphs in the cell
        paragraphs = tc_element.findall(qn('w:p'))
        if not paragraphs:
            return

        # Use first paragraph, remove extras
        first_p = paragraphs[0]
        for extra_p in paragraphs[1:]:
            tc_element.remove(extra_p)

        # Clear existing runs
        for run in list(first_p.findall(qn('w:r'))):
            first_p.remove(run)

        # Add new run with text
        run_el = first_p.makeelement(qn('w:r'), {})

        # Copy run properties from template if available
        # (preserves font, size, color)
        old_runs = first_p.findall(qn('w:r'))
        # No old runs left, but we can check pPr for inherited style

        t_el = run_el.makeelement(qn('w:t'), {})
        t_el.text = text
        t_el.set(qn('xml:space'), 'preserve')
        run_el.append(t_el)
        first_p.append(run_el)
