"""
Voorlees Mededelingen Generator
Generates a 2-column Word document for reading aloud during service.
Left column = Dutch (black), Right column = Indonesian (dark blue).

Sections:
  1. Welkomswoord       – full text from mededelingen, translated
  2. Collecte Opbrengsten – single merged cell with embedded table
  3. Dankoffer           – as in mededelingen but without URL, with QR
  4. Regionale Mededelingen – each item (heading+body) in its own row
  5. Landelijke Mededelingen – same as regionale
"""

import os
import re
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import Config

ID_COLOR  = RGBColor(0x00, 0x35, 0x7A)
NL_COLOR  = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = 'Aptos Narrow'

DUTCH_MONTHS = ['januari','februari','maart','april','mei','juni',
                'juli','augustus','september','oktober','november','december']
DUTCH_DAYS   = ['maandag','dinsdag','woensdag','donderdag','vrijdag','zaterdag','zondag']
INDO_MONTHS  = ['Januari','Februari','Maret','April','Mei','Juni',
                'Juli','Agustus','September','Oktober','November','Desember']
INDO_DAYS    = ['Senin','Selasa','Rabu','Kamis','Jumat','Sabtu','Minggu']


def _format_nl_date(dt: datetime) -> str:
    return f"{dt.day} {DUTCH_MONTHS[dt.month - 1]} {dt.year}"


def _format_id_date(dt: datetime) -> str:
    return f"{dt.day} {INDO_MONTHS[dt.month - 1]} {dt.year}"


# ── Low-level helpers ────────────────────────────────────────────────────────

def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _shade_cell(cell, fill='D9D9D9'):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _set_col_width(cell, cm_val):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(Cm(cm_val).twips)))
    tcW.set(qn('w:type'), 'dxa')


def _r(para, text: str, bold=False, color=NL_COLOR, size_pt=12):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.color.rgb = color
    run.font.size = Pt(size_pt)
    return run


def _section_header_row(table, nl_text: str, id_text: str):
    """Blue merged row spanning both columns — section title."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    _set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
    _set_col_width(cell, 16.0)   # full width
    _shade_cell(cell, 'BDD7EE')
    p = cell.paragraphs[0]
    _r(p, nl_text, bold=True, color=NL_COLOR, size_pt=14)
    _r(p, '  /  ', bold=True, color=RGBColor(0x44, 0x44, 0x44), size_pt=14)
    _r(p, id_text, bold=True, color=ID_COLOR, size_pt=14)


def _content_row(table, nl_parts: list, id_parts: list):
    """
    Add one content row.
    Each part is (text, bold).  Parts in the same cell are written consecutively
    on newlines when the previous part ended with \\n.
    """
    row = table.add_row()
    nl_cell = row.cells[0]
    id_cell = row.cells[1]
    _set_cell_margins(nl_cell)
    _set_cell_margins(id_cell)
    _set_col_width(nl_cell, 8.0)
    _set_col_width(id_cell, 8.0)

    def fill_cell(cell, parts, base_color):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        first = True
        for text, bold in parts:
            if not first:
                p.add_run('\n')
            color = base_color
            _r(p, text, bold=bold, color=color, size_pt=12)
            first = False

    fill_cell(nl_cell, nl_parts, NL_COLOR)
    fill_cell(id_cell, id_parts, ID_COLOR)
    return row


def _subheader_row(table, nl_text: str, id_text: str):
    """Bold sub-header row (white background) within a section."""
    _content_row(table,
                 [(nl_text, True)],
                 [(id_text, True)])


def _parse_meded_blocks(text: str) -> List[dict]:
    """
    Split raw mededelingen text into blocks of {'heading': str, 'body': str}.
    Blank lines separate blocks. First line of each block is the heading if it
    is short (≤80 chars) and does not end in sentence-ending punctuation.
    Nothing is ever dropped.
    """
    if not text.strip():
        return []
    blocks = []
    # Split on blank lines to get paragraph groups
    raw_blocks = re.split(r'\n\s*\n', text.replace('\r', ''))
    for block in raw_blocks:
        lines = [l.strip() for l in block.split('\n') if l.strip()]
        if not lines:
            continue
        first = lines[0]
        is_heading = (len(first) <= 80 and
                      not first[-1] in ('.', ',', ';', '!') if first else False)
        if is_heading and len(lines) > 1:
            blocks.append({'heading': first, 'body': '\n'.join(lines[1:])})
        else:
            blocks.append({'heading': '', 'body': '\n'.join(lines)})
    return blocks


def _align_id_blocks(nl_blocks: list, id_blocks: list) -> list:
    """
    Ensure id_blocks has the same length as nl_blocks by merging extra
    consecutive ID blocks (those without a heading) into the preceding block.
    This handles cases where the ID translation has extra blank-line splits.
    """
    if len(id_blocks) <= len(nl_blocks):
        # Pad with empty blocks if ID has fewer
        result = list(id_blocks)
        while len(result) < len(nl_blocks):
            result.append({'heading': '', 'body': ''})
        return result

    # Merge forward: collapse headingless blocks into previous
    merged = []
    for blk in id_blocks:
        if blk['heading'] == '' and merged:
            prev = merged[-1]
            sep = '\n' if prev['body'] else ''
            prev['body'] = prev['body'] + sep + blk['body']
        else:
            merged.append({'heading': blk['heading'], 'body': blk['body']})
    # If still more than nl_blocks, merge tail into last
    while len(merged) > len(nl_blocks) and len(merged) > 1:
        last = merged.pop()
        sep = '\n' if merged[-1]['body'] else ''
        merged[-1]['body'] = merged[-1]['body'] + sep + last['heading'] + ('\n' if last['heading'] and last['body'] else '') + last['body']
    # Pad if needed
    while len(merged) < len(nl_blocks):
        merged.append({'heading': '', 'body': ''})
    return merged


# ── Generator class ───────────────────────────────────────────────────────────

class VoorleesGenerator:
    def __init__(self):
        self.output_dir = Config.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    def _write_with_bold_names(self, para, text: str, names: list, color):
        """Write text into para, making each occurrence of any name in names bold."""
        import re
        # Build combined pattern for all non-empty names
        names = [n for n in names if n]
        if not names:
            _r(para, text, color=color)
            return
        pattern = '|'.join(re.escape(n) for n in sorted(names, key=len, reverse=True))
        parts = re.split(f'({pattern})', text)
        for part in parts:
            if not part:
                continue
            bold = part in names
            _r(para, part, bold=bold, color=color)

    def _build_id_welkom_segments(self, nl_lines: list, predikant: str, ovd: str, is_ole: bool, opmerking: str,
                                   day_id: str, date_str_id: str) -> list:
        """
        Returns a list of lines, each line is a list of (text, bold) segments.
        Uses hardcoded Indonesian template with bold names injected.
        """
        import re
        result = []

        def seg(text, bold=False): return (text, bold)

        def split_bold(text, *names):
            """Split text into (text, bold) segments around given names."""
            parts_out = []
            remaining = text
            for name in names:
                if name and name in remaining:
                    idx = remaining.index(name)
                    if idx > 0:
                        parts_out.append(seg(remaining[:idx]))
                    parts_out.append(seg(name, bold=True))
                    remaining = remaining[idx + len(name):]
            if remaining:
                parts_out.append(seg(remaining))
            return parts_out

        # Extract dienst type from opmerking
        dienst_type = ''
        if opmerking:
            text = opmerking.split('OLE')[0].strip().rstrip(',').strip()
            if text:
                dienst_type = text + ' '
        
        # Translate dienst type to Indonesian
        dienst_type_id = ''
        if dienst_type:
            dienst_lower = dienst_type.lower()
            if 'pinksteren' in dienst_lower:
                dienst_type_id = 'Pentakosta '
            elif 'hemelvaart' in dienst_lower:
                dienst_type_id = 'Kenaikan Yesus Kristus '
            else:
                dienst_type_id = dienst_type  # Keep as-is for unknown types

        print(f"DEBUG _build_id_welkom_segments: opmerking='{opmerking}', dienst_type='{dienst_type}', dienst_type_id='{dienst_type_id}'")
        print(f"DEBUG _build_id_welkom_segments: Processing {len(nl_lines)} lines:")
        for i, line in enumerate(nl_lines):
            print(f"  {i}: '{line}'")

        for s in nl_lines:
            s = s.strip()
            if not s:
                continue
            if s.startswith('Goedemorgen'):
                result.append([seg('Selamat pagi saudara-saudari,')])
            elif 'van harte welkom' in s or s.startswith('Namens de kerkenraad'):
                # Dynamic Online/Offline and dienst type handling
                # Correct order: "ibadah [online] [dienst_type] ini"
                online_text = 'online ' if is_ole else ''
                dienst_text = f'{online_text}{dienst_type_id}'.strip()
                if dienst_text:
                    dienst_text += ' '
                result.append([seg(
                    f'Atas nama Majelis regio Amstelveen, saya mengucapkan selamat datang '
                    f'dalam {dienst_text}ibadah ini, khususnya bagi mereka yang baru pertama kali hadir.'
                )])
            elif s.startswith('Vandaag,'):
                line_segs = (
                    [seg('Hari ini, ibadah akan dipimpin oleh ')] +
                    [seg(predikant, True)] +
                    [seg('. Majelis yang bertugas dalam ibadah adalah ')] +
                    [seg(ovd, True)] +
                    [seg('. Jika anda memiliki pertanyaan, dapat menghubungi beliau.')]
                )
                result.append(line_segs)
            elif 'aanstaande' in s.lower():
                m_date = re.search(r'(\d{1,2}\s+\w+\s+\d{4})', s)
                m_time = re.search(r'(\d{1,2}[:.\u00b7]\d{2})\s*uur', s)
                
                # Extract dienst type and online status using simple keyword search
                s_lower = s.lower()
                is_online = 'online' in s_lower
                dienst_type = ''
                if 'pinksteren' in s_lower:
                    dienst_type = 'Pinksteren'
                elif 'hemelvaart' in s_lower:
                    dienst_type = 'Hemelvaart'
                
                print(f"DEBUG aanstaande: s='{s}', is_online={is_online}, dienst_type='{dienst_type}'")
                # Match name after 'voor te gaan,' up to '. Aanvang' — allow dots inside (ds., pdt.)
                m_pred = re.search(r'voor te gaan,?\s+(.+?)(?=\.\s*Aanvang|\. Aanvang|$)', s, re.IGNORECASE)
                date_part = m_date.group(1) if m_date else ''
                for nl_m, id_m in zip(DUTCH_MONTHS, INDO_MONTHS):
                    date_part = re.sub(rf'\b{nl_m}\b', id_m, date_part, flags=re.IGNORECASE)
                time_part = m_time.group(1).replace('.', ':') if m_time else '10:30'
                pred_part = m_pred.group(1).strip().rstrip('. ') if m_pred else ''
                
                # Translate dienst type to Indonesian
                dienst_type_id = ''
                if dienst_type:
                    dienst_lower = dienst_type.lower()
                    if 'pinksteren' in dienst_lower:
                        dienst_type_id = 'Pentakosta '
                    elif 'hemelvaart' in dienst_lower:
                        dienst_type_id = 'Kenaikan Yesus Kristus '
                    else:
                        dienst_type_id = dienst_type + ' '
                
                print(f"DEBUG aanstaande: dienst_type_id='{dienst_type_id}'")
                
                if 'donderdag' in s.lower() or 'hemelvaart' in s.lower():
                    dienst = 'ibadah Kebangkitan Yesus Kristus'
                    day_word = 'Kamis'
                else:
                    dienst = 'ibadah'
                    day_word = 'Minggu'
                
                # Build Indonesian text: "ibadah [online] [dienst_type]" to match first paragraph pattern
                online_text = 'Online ' if is_online else ''
                line_segs = (
                    [seg(f'Pada {day_word} yang akan datang, {date_part}, ibadah {online_text}{dienst_type_id}di Amstelveen akan dipimpin oleh ')] +
                    [seg(pred_part, True)] +
                    [seg(f'. Kebaktian akan dimulai pada pukul {time_part}.')]
                )
                result.append(line_segs)
            else:
                result.append([seg(s)])
        return result

    def _build_opbrengst_table(self, doc, entries, flat):
        """Build a standalone opbrengst table with matching blue section header."""
        def fmtv(v): return v.strip() if v else '–'
        def parse_eur(s):
            s = (s or '').replace('€','').replace('.','').replace(',','.').strip()
            try: return float(s)
            except: return 0.0

        datasets = entries if entries else ([flat] if flat.get('collecte_contant') else [])

        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'

        # Section header row — same blue style as main table headers
        hrow = t.add_row()
        hrow.cells[0].merge(hrow.cells[1])
        hcell = hrow.cells[0]
        _set_cell_margins(hcell, top=100, start=100, bottom=100, end=100)
        _set_col_width(hcell, 16.0)
        _shade_cell(hcell, 'BDD7EE')
        hp = hcell.paragraphs[0]
        _r(hp, 'Collecte Opbrengsten', bold=True, color=NL_COLOR, size_pt=14)
        _r(hp, '  /  ', bold=True, color=RGBColor(0x44, 0x44, 0x44), size_pt=14)
        _r(hp, 'Hasil Kolekte', bold=True, color=ID_COLOR, size_pt=14)

        if not datasets:
            er = t.add_row()
            er.cells[0].merge(er.cells[1])
            _r(er.cells[0].paragraphs[0], '(geen gegevens)', color=NL_COLOR)
            return t

        def add_row(label, amount='', bold_label=False, bold_amount=False, span=False):
            row = t.add_row()
            c0, c1 = row.cells[0], row.cells[1]
            _set_col_width(c0, 8.0)
            _set_col_width(c1, 8.0)
            _set_cell_margins(c0, top=40, start=80, bottom=40, end=40)
            _set_cell_margins(c1, top=40, start=40, bottom=40, end=80)
            if span:
                c0.merge(c1)
                _r(c0.paragraphs[0], label, bold=bold_label, color=NL_COLOR, size_pt=12)
            else:
                _r(c0.paragraphs[0], label, bold=bold_label, color=NL_COLOR, size_pt=12)
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _r(p1, amount, bold=bold_amount, color=NL_COLOR, size_pt=12)

        for idx, e in enumerate(datasets):
            lbl     = e.get('date_label', '')
            contant = fmtv(e.get('collecte_contant', ''))
            bonnen  = fmtv(e.get('collecte_bonnen', ''))
            bank    = fmtv(e.get('collecte_bank', ''))
            tikkie  = fmtv(e.get('collecte_tikkie', ''))
            ole     = fmtv(e.get('collecte_ole', ''))
            bv      = fmtv(e.get('bezoekers_volwassenen', ''))
            bk      = fmtv(e.get('bezoekers_kinderen', ''))
            try:    bt = str(int(bv) + int(bk))
            except: bt = ''

            vals    = [e.get(k,'') for k in
                       ['collecte_contant','collecte_bonnen','collecte_bank','collecte_tikkie']]
            total_f = sum(parse_eur(v) for v in vals)
            # Format as Dutch: 810,63
            total_s = f'{total_f:,.2f}'.replace(',','X').replace('.',',').replace('X','.')

            if idx > 0:
                add_row('', span=True)  # spacer between entries

            if lbl:           add_row(lbl, bold_label=True, span=True)
            add_row('Collecte', bold_label=True, span=True)
            if contant != '–': add_row('1. Contant',          f'€  {contant}')
            if bonnen  != '–': add_row('2. Collectebonnen',   f'€  {bonnen}')
            if bank    != '–': add_row('3. Bankoverschrijving',f'€  {bank}')
            if tikkie  != '–': add_row('4. TIKKIE',           f'€  {tikkie}')
            add_row('Totaal collecte opbrengst:', f'€  {total_s}',
                    bold_label=True, bold_amount=True)
            if bv != '–' or bk != '–':
                vis = f'Aantal bezoekers: {bv} volwassenen, {bk} kinderen'
                if bt: vis += f', {bt} totaal bezoekers'
                add_row(vis, span=True)
            add_row('', span=True)  # blank spacer
            if ole != '–':
                add_row(f'Collecte opbrengst OLE {lbl}', f'€  {ole}')

        return t

    def generate(self, mededelingen_date: datetime,
                 takenrooster_entry: Dict[str, Any],
                 mededelingen_data: Dict[str, Any],
                 user_data: Dict[str, Any],
                 welkom_paragraphs: List[str] = None) -> str:

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(1.5)
            sec.bottom_margin = Cm(1.5)
            sec.left_margin   = Cm(1.8)
            sec.right_margin  = Cm(1.8)

        predikant = takenrooster_entry.get('predikant', '')
        ovd       = takenrooster_entry.get('ovd', '')
        opmerking = takenrooster_entry.get('opmerking', '')
        is_ole    = 'OLE' in opmerking.upper()
        eo1       = takenrooster_entry.get('1eo', '')
        beamer    = takenrooster_entry.get('beamer', '')
        date_str_nl = _format_nl_date(mededelingen_date)
        date_str_id = _format_id_date(mededelingen_date)
        day_nl = DUTCH_DAYS[mededelingen_date.weekday()]
        day_id = INDO_DAYS[mededelingen_date.weekday()]

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(title, f"Vertaling Mededelingen – {_format_nl_date(mededelingen_date)}",
           bold=True, color=NL_COLOR, size_pt=16)

        # OvD / 1eO / Beamer info block
        info = doc.add_paragraph()
        _r(info, 'OvD: ', bold=True, color=NL_COLOR, size_pt=11)
        _r(info, ovd or '–', color=NL_COLOR, size_pt=11)
        if eo1:
            _r(info, '   |   1eO: ', bold=True, color=NL_COLOR, size_pt=11)
            _r(info, eo1, color=NL_COLOR, size_pt=11)
        if beamer:
            _r(info, '   |   Beamer: ', bold=True, color=NL_COLOR, size_pt=11)
            _r(info, beamer, color=NL_COLOR, size_pt=11)

        # Column label
        hdr = doc.add_paragraph()
        _r(hdr, 'Nederlands', bold=True, color=NL_COLOR, size_pt=11)
        _r(hdr, '    |    ', color=RGBColor(0x88, 0x88, 0x88), size_pt=11)
        _r(hdr, 'Indonesia', bold=True, color=ID_COLOR, size_pt=11)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # ── 1. Welkomswoord ──────────────────────────────────────────────
        _section_header_row(table, 'Welkomswoord', 'Kata Sambutan')

        if welkom_paragraphs:
            nl_welkom = '\n'.join(p for p in welkom_paragraphs if p.strip())
        else:
            nl_welkom = (
                f"Goedemorgen broeders en zusters,\n"
                f"Namens de kerkenraad heet ik u allen van harte welkom bij deze eredienst, "
                f"in het bijzonder diegenen die vandaag voor het eerst aanwezig zijn.\n"
                f"Vandaag, {day_nl} {date_str_nl}, gaat voor {predikant}. "
                f"De ouderling van dienst is {ovd}. "
                f"Als u vragen heeft, kunt u de ouderling van dienst aanspreken."
            )

        nl_lines = nl_welkom.split('\n')
        # id_segments: list of (text, bold) tuples per line, newline-separated
        id_segments = self._build_id_welkom_segments(nl_lines, predikant, ovd, is_ole, opmerking, day_id, date_str_id)

        # Welkom row — NL left (black), ID right (dark blue), names bold
        row = table.add_row()
        nl_cell = row.cells[0]; id_cell = row.cells[1]
        _set_cell_margins(nl_cell); _set_cell_margins(id_cell)
        _set_col_width(nl_cell, 8.0); _set_col_width(id_cell, 8.0)

        # Fill NL cell — bold predikant and ovd wherever they appear
        p_nl = nl_cell.paragraphs[0]
        for i, line in enumerate(nl_lines):
            if i: p_nl.add_run('\n')
            self._write_with_bold_names(p_nl, line.strip(), [predikant, ovd], NL_COLOR)

        # Fill ID cell — bold predikant and ovd wherever they appear
        p_id = id_cell.paragraphs[0]
        first_seg = True
        for seg_line in id_segments:   # each element is list of (text, bold)
            if not first_seg: p_id.add_run('\n')
            first_seg = False
            for txt, bold in seg_line:
                _r(p_id, txt, bold=bold, color=ID_COLOR)

        # ── End of table 1 (Welkomswoord) ─────────────────────────────
        # ── 2. Collecte Opbrengsten — table with matching header style ───
        opbrengst_entries = user_data.get('opbrengst_entries', [])
        self._build_opbrengst_table(doc, opbrengst_entries, user_data)

        # ── 3. Dankoffer + Mededelingen — new 2-column table ─────────────
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _section_header_row(table, 'Dankoffer', 'Persembahan Syukur')

        dankoffer_qr = user_data.get('dankoffer_qr', '')
        ole_qr       = user_data.get('ole_qr', '')
        tikkie_url   = user_data.get('dankoffer_url', '')
        ole_url      = user_data.get('ole_url', '')
        date_ovv     = mededelingen_date.strftime('%d-%m-%Y')

        # NL dankoffer text (matches mededelingen layout, URL removed)
        nl_dank = (
            f"Het dankoffer voor het werk van de kerk kunt u tijdens deze eredienst "
            f"geven op drie manieren:\n"
            f"1.  Deponeren in de collectezak\n"
            f"2.  Overmaken naar rekeningnummer van GKIN Amstelveen, "
            f"IBAN: NL40.ABNA.0549.3085.12, o.v.v. \"Collecte {date_ovv}\"\n"
            f"3.  Gebruik te maken van de QR-code, of betaalverzoek link:"
        )
        id_dank = (
            f"Persembahan untuk pelayanan gereja dapat anda berikan dalam ibadah ini dengan:\n"
            f"1.  Memasukkan dalam kantong persembahan yang ada di gereja.\n"
            f"2.  Transfer ke rekening GKIN Amstelveen, IBAN: NL40.ABNA.0549.3085.12, "
            f"o.v.v. \"Collecte {date_ovv}\"\n"
            f"3.  Menggunakan QR-code yang ditampilkan di layar:"
        )

        row = table.add_row()
        nl_cell = row.cells[0]
        id_cell = row.cells[1]
        _set_cell_margins(nl_cell)
        _set_cell_margins(id_cell)
        _set_col_width(nl_cell, 8.0)
        _set_col_width(id_cell, 8.0)

        # NL cell: text then QR image
        p_nl = nl_cell.paragraphs[0]
        _r(p_nl, nl_dank, color=NL_COLOR, size_pt=12)
        if dankoffer_qr and os.path.exists(dankoffer_qr):
            try:
                p_nl.add_run('\n')
                p_nl.add_run().add_picture(dankoffer_qr, width=Cm(3.5))
            except Exception as e:
                _r(p_nl, f'\n[QR: {e}]', color=NL_COLOR, size_pt=9)

        # ID cell: text then dankoffer QR again (same QR shown in both columns)
        p_id = id_cell.paragraphs[0]
        _r(p_id, id_dank, color=ID_COLOR, size_pt=12)
        if dankoffer_qr and os.path.exists(dankoffer_qr):
            try:
                p_id.add_run('\n')
                p_id.add_run().add_picture(dankoffer_qr, width=Cm(3.5))
            except Exception:
                pass

        # ── 4. Regionale Mededelingen ─────────────────────────────────────
        _section_header_row(table, 'Regionale Mededelingen', 'Berita Regional')
        reg_nl = mededelingen_data.get('regionale_nl', '').strip()
        reg_id = mededelingen_data.get('regionale_id', '').strip()
        nl_blocks = _parse_meded_blocks(reg_nl)
        id_blocks = _align_id_blocks(nl_blocks, _parse_meded_blocks(reg_id))
        max_b = max(len(nl_blocks), len(id_blocks))
        for i in range(max_b):
            nb = nl_blocks[i] if i < len(nl_blocks) else {'heading': '', 'body': ''}
            ib = id_blocks[i] if i < len(id_blocks) else {'heading': '', 'body': ''}
            nl_parts = []
            id_parts = []
            if nb['heading']:
                nl_parts.append((nb['heading'], True))
            if nb['body']:
                nl_parts.append((nb['body'], False))
            if not nl_parts:
                nl_parts = [('', False)]
            if ib['heading']:
                id_parts.append((ib['heading'], True))
            if ib['body']:
                id_parts.append((ib['body'], False))
            if not id_parts:
                id_parts = [('', False)]
            _content_row(table, nl_parts, id_parts)

        # ── 5. Landelijke Mededelingen ────────────────────────────────────
        _section_header_row(table, 'Landelijke Mededelingen', 'Berita Nasional')
        land_nl = mededelingen_data.get('landelijke_nl', '').strip()
        land_id = mededelingen_data.get('landelijke_id', '').strip()
        nl_blocks = _parse_meded_blocks(land_nl)
        id_blocks = _align_id_blocks(nl_blocks, _parse_meded_blocks(land_id))
        max_b = max(len(nl_blocks), len(id_blocks))
        for i in range(max_b):
            nb = nl_blocks[i] if i < len(nl_blocks) else {'heading': '', 'body': ''}
            ib = id_blocks[i] if i < len(id_blocks) else {'heading': '', 'body': ''}
            nl_parts = []
            id_parts = []
            if nb['heading']:
                nl_parts.append((nb['heading'], True))
            if nb['body']:
                nl_parts.append((nb['body'], False))
            if not nl_parts:
                nl_parts = [('', False)]
            if ib['heading']:
                id_parts.append((ib['heading'], True))
            if ib['body']:
                id_parts.append((ib['body'], False))
            if not id_parts:
                id_parts = [('', False)]
            _content_row(table, nl_parts, id_parts)

        # Save
        fname    = f'Vertaling_Mededelingen_{mededelingen_date.strftime("%y%m%d")}.docx'
        filepath = os.path.join(self.output_dir, fname)
        doc.save(filepath)
        print(f"Voorlees doc saved: {filepath}")
        return filepath
