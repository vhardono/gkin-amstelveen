"""
Preekbevestiging Document Generator
Copies the template docx and fills in dynamic values using regex-based
substitution — works regardless of which prior values are in the template.
"""

import os
import re
from datetime import datetime, timedelta

from docx import Document

DUTCH_MONTHS = ['januari','februari','maart','april','mei','juni',
                'juli','augustus','september','oktober','november','december']
DUTCH_DAYS   = ['maandag','dinsdag','woensdag','donderdag',
                'vrijdag','zaterdag','zondag']

_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(_DIR, 'doc_templates', 'preekbevestiging_template.docx')


def _fmt_date_long(iso_str: str) -> str:
    d = datetime.strptime(iso_str, '%Y-%m-%d')
    return f"Zondag {d.day} {DUTCH_MONTHS[d.month-1]} {d.year}"


def _fmt_date_today() -> str:
    d = datetime.now()
    return f"Amstelveen, {d.day} {DUTCH_MONTHS[d.month-1]} {d.year}"


def _deadline_tue(iso_str: str) -> str:
    d = datetime.strptime(iso_str, '%Y-%m-%d')
    tue = d - timedelta(days=5)
    return f"dinsdag {tue.day} {DUTCH_MONTHS[tue.month-1]} {tue.year}"


# ---------------------------------------------------------------------------
# Regex-based replacers — match any existing value, replace with new one
# ---------------------------------------------------------------------------

# Dutch date pattern: "Zondag 10 mei 2026"
_RE_DATE_LONG = re.compile(
    r'Zondag\s+\d{1,2}\s+(?:januari|februari|maart|april|mei|juni|'
    r'juli|augustus|september|oktober|november|december)\s+\d{4}'
)
# "Amstelveen, 9 mei 2026"
_RE_DATE_TODAY = re.compile(
    r'Amstelveen,\s+\d{1,2}\s+(?:januari|februari|maart|april|mei|juni|'
    r'juli|augustus|september|oktober|november|december)\s+\d{4}'
)
# deadline: "dinsdag 5 mei 2026"
_RE_DEADLINE = re.compile(
    r'dinsdag\s+\d{1,2}\s+(?:januari|februari|maart|april|mei|juni|'
    r'juli|augustus|september|oktober|november|december)\s+\d{4}'
)
# service date + time in body: "Zondag 10 mei 2026 om 10:30 uur"
_RE_DATE_TIME = re.compile(
    r'Zondag\s+\d{1,2}\s+\S+\s+\d{4}\s+om\s+\d{1,2}[:.]\d{2}\s+uur'
)
# time cell: "10:30 uur"
_RE_TIME_CELL = re.compile(r'^\d{1,2}[:.]\d{2}\s+uur$')


def _sub_para(para, pattern, replacement: str):
    """Apply regex substitution to a paragraph, preserving first-run formatting."""
    full = para.text
    if not pattern.search(full):
        return False
    new_text = pattern.sub(replacement, full)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ''
    return True


def _sub_cell(cell, pattern, replacement: str):
    for para in cell.paragraphs:
        _sub_para(para, pattern, replacement)


def _set_cell_text(cell, new_text: str):
    """Replace all text in a cell's first paragraph, preserving run formatting."""
    for para in cell.paragraphs:
        if para.text.strip():
            for i, run in enumerate(para.runs):
                run.text = new_text if i == 0 else ''
            if not para.runs:
                para.add_run(new_text)
            return


def generate(entry: dict, iso_date: str, output_path: str, songs: list = None):
    doc = Document(TEMPLATE_PATH)

    predikant  = entry.get('predikant', '')
    ovd        = entry.get('ovd', '')
    eo1        = entry.get('1eo', '') or entry.get('eo1', '')
    tijd       = (entry.get('tijd', '') or '10:30').replace('.', ':')
    if ':' not in tijd:
        tijd = '10:30'

    parts = predikant.strip().split()
    title = parts[0] if parts else ''
    last  = parts[-1] if len(parts) > 1 else predikant
    last_cap = last[0].upper() + last[1:] if last else last
    greeting_name = f"{title} {last_cap}".strip()

    date_long  = _fmt_date_long(iso_date)   # "Zondag 17 mei 2026"
    date_today = _fmt_date_today()           # "Amstelveen, 15 mei 2026"
    deadline   = _deadline_tue(iso_date)     # "dinsdag 12 mei 2026"
    date_time  = f"{date_long} om {tijd} uur"

    # ------------------------------------------------------------------ #
    # Paragraphs
    # ------------------------------------------------------------------ #
    for para in doc.paragraphs:
        txt = para.text
        # "Amstelveen, <date>"
        _sub_para(para, _RE_DATE_TODAY, date_today)
        # "Zondag X Y Z om HH:MM uur" (body service line)
        _sub_para(para, _RE_DATE_TIME, date_time)
        # standalone "Zondag X Y Z" not already handled above
        _sub_para(para, _RE_DATE_LONG, date_long)
        # deadline lines
        if 'uiterlijk' in txt and 'invulformulier' in txt:
            _sub_para(para, _RE_DEADLINE, deadline)
        elif 'uiterlijk' in txt and txt.rstrip().endswith(')'):
            _sub_para(para, _RE_DEADLINE, deadline)
        # Addressee line (bare predikant name)
        if re.match(r'^(?:ds\.|zr\.|br\.)\s+\S', txt.strip()):
            for i, run in enumerate(para.runs):
                run.text = predikant if i == 0 else ''
        # Salutation
        if txt.strip().startswith('Geachte'):
            for i, run in enumerate(para.runs):
                run.text = f'Geachte {greeting_name},' if i == 0 else ''

    # ------------------------------------------------------------------ #
    # Table cells
    # ------------------------------------------------------------------ #
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                # Date+time body line
                _sub_cell(cell, _RE_DATE_TIME, date_time)
                # Standalone date
                _sub_cell(cell, _RE_DATE_LONG, date_long)
                # Today's date header
                _sub_cell(cell, _RE_DATE_TODAY, date_today)
                # Time cell (e.g. "10:30 uur")
                if _RE_TIME_CELL.match(txt):
                    _set_cell_text(cell, f'{tijd} uur')
                # Predikant cell — matches "ds./zr./br. ..."
                elif re.match(r'^(?:ds\.|zr\.|br\.)\s+\S', txt):
                    _set_cell_text(cell, predikant)
                # OvD cell — any non-predikant name row following "Ouderling van Dienst"
                elif ovd and txt and txt == cell.text.strip() and _looks_like_name(txt) and txt != predikant:
                    pass  # handled by row position below

    # OvD and 1e Ontvangst — positional: Table 1, rows 4 and 5, col 2
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        rows = t1.rows
        # Row 4 (0-indexed) = Ouderling van Dienst value
        if len(rows) > 4 and len(rows[4].cells) > 2 and ovd:
            _set_cell_text(rows[4].cells[2], ovd)
        # Row 5 = Eerste Ontvangst value
        if len(rows) > 5 and len(rows[5].cells) > 2 and eo1:
            _set_cell_text(rows[5].cells[2], eo1)

    # ------------------------------------------------------------------ #
    # Song table (Table 2) — col 2: always clear first, then write value
    # ------------------------------------------------------------------ #
    if len(doc.tables) > 2:
        song_table = doc.tables[2]
        for ri, row in enumerate(song_table.rows):
            cells = row.cells
            if len(cells) > 2:
                new_val = (songs[ri].strip() if songs and ri < len(songs) and songs[ri] else '……………………………')
                cell = cells[2]
                # Clear all runs in all paragraphs
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                # Write new value into first paragraph's first run (or add run)
                first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                if first_para.runs:
                    first_para.runs[0].text = new_val
                else:
                    first_para.add_run(new_val)

    doc.save(output_path)


def _looks_like_name(txt: str) -> bool:
    """Heuristic: at least two capitalised words, no digits."""
    if re.search(r'\d', txt):
        return False
    words = txt.split()
    return len(words) >= 2 and sum(1 for w in words if w[0].isupper()) >= 1
